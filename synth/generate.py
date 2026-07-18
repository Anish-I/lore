#!/usr/bin/env python3
"""Synthetic corpus generator for the enterprise work-tool.

Materializes two coherent, multi-format "people-work" scenarios that the
wizard/plugin layer is built and evaluated against:

  * enterprise/  — a Q3 budget-overrun + hiring-freeze thread (broad appeal)
  * municipal/   — a public-records (FOIA) request thread (the vertical slice)

Each scenario is an email chain PLUS the documents it references (docx / xlsx /
pdf) so the full ingestion path (PyMuPDF for PDF, the zip/XML DOCX reader,
openpyxl-shaped sheets) gets exercised end-to-end, and a ground-truth to-do
list (synth/expected/*.json) the "email chain -> to-dos" wizard is scored on.

No real data exists yet (per Anish, 2026-07-17) — this is the synthetic stand-in.
Regenerate with:  python synth/generate.py
Output lands in synth/out/ (gitignored — a build artifact, not source).

Deps: python-docx, openpyxl, reportlab (all present on the dev box). pptx is
unavailable, so "presentations" are rendered as PDF slide-decks (still ingestible).
"""
from __future__ import annotations

import json
import os
import textwrap

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "out")
EXPECTED = os.path.join(HERE, "expected")

# --- Scenario content (data, kept separate from rendering) -------------------
# Absolute dates anchored to the handover context (today = 2026-07-17).

ENTERPRISE_THREAD = {
    "filename": "enterprise/emails/q3-budget-thread.eml",
    "subject": "Re: Q3 Budget Review — action items",
    # newest message first; older ones quoted below (as real exports look).
    "messages": [
        {
            "from": "Dana Ruiz <dana.ruiz@northwind.example>",
            "to": "Marcus Bell <marcus.bell@northwind.example>, "
                  "Priya Nair <priya.nair@northwind.example>, "
                  "Mahmoud Hafez <mahmoud.hafez@northwind.example>",
            "date": "Wed, 15 Jul 2026 17:42:00 -0400",
            "body": textwrap.dedent("""\
                Thanks all. To close this out before the board meeting:

                - Marcus, send me the revised engineering headcount plan by Friday EOD.
                - Priya, please pause the two open backend reqs until we reforecast.
                - Mahmoud, book the finance review room for Thursday 2pm and circulate
                  the deck beforehand.

                I'll update the board on Monday.

                Dana"""),
        },
        {
            "from": "Marcus Bell <marcus.bell@northwind.example>",
            "to": "Dana Ruiz <dana.ruiz@northwind.example>",
            "date": "Wed, 15 Jul 2026 16:10:00 -0400",
            "body": textwrap.dedent("""\
                The variance is mostly cloud spend — full breakdown in
                q3-budget.xlsx. I can pull about $40k out of tooling this quarter
                without hurting delivery."""),
        },
        {
            "from": "Priya Nair <priya.nair@northwind.example>",
            "to": "Dana Ruiz <dana.ruiz@northwind.example>",
            "date": "Wed, 15 Jul 2026 15:55:00 -0400",
            "body": textwrap.dedent("""\
                HR note: per spend-freeze-memo.docx, all non-critical hiring is
                paused this quarter. The two backend reqs are not on the critical
                list, so we can hold them."""),
        },
        {
            "from": "Dana Ruiz <dana.ruiz@northwind.example>",
            "to": "Marcus Bell <marcus.bell@northwind.example>, "
                  "Priya Nair <priya.nair@northwind.example>, "
                  "Mahmoud Hafez <mahmoud.hafez@northwind.example>",
            "date": "Wed, 15 Jul 2026 14:30:00 -0400",
            "body": textwrap.dedent("""\
                Team, we're tracking 12% over on Q3 opex. I need a concrete plan
                before the board meeting. Deck attached (q3-review.pdf); numbers
                in q3-budget.xlsx.

                Dana Ruiz
                CFO, Northwind"""),
        },
    ],
}

ENTERPRISE_TODOS = [
    {"assignee": "Marcus Bell", "task": "Send revised engineering headcount plan",
     "due": "2026-07-17", "due_text": "Friday EOD", "source": "Dana Ruiz, 15 Jul 17:42"},
    {"assignee": "Priya Nair", "task": "Pause the two open backend reqs until reforecast",
     "due": None, "due_text": None, "source": "Dana Ruiz, 15 Jul 17:42"},
    {"assignee": "Mahmoud Hafez", "task": "Book the finance review room for Thursday 2pm",
     "due": "2026-07-16", "due_text": "Thursday 2pm", "source": "Dana Ruiz, 15 Jul 17:42"},
    {"assignee": "Mahmoud Hafez", "task": "Circulate the Q3 review deck before Thursday",
     "due": "2026-07-16", "due_text": "before Thursday", "source": "Dana Ruiz, 15 Jul 17:42"},
    {"assignee": "Dana Ruiz", "task": "Update the board on the Q3 opex plan",
     "due": "2026-07-20", "due_text": "Monday", "source": "Dana Ruiz, 15 Jul 17:42"},
]

MUNICIPAL_THREAD = {
    "filename": "municipal/emails/foia-2026-0417-thread.eml",
    "subject": "FOIA #2026-0417 — police overtime records (Jan–Jun)",
    "messages": [
        {
            "from": "Sofia Marin <clerk@cityofrivermont.example>",
            "to": "Alan Woods <legal@cityofrivermont.example>, "
                  "Mahmoud Hafez <records@cityofrivermont.example>",
            "date": "Thu, 16 Jul 2026 09:12:00 -0500",
            "body": textwrap.dedent("""\
                Public-records request #2026-0417 received from J. Alvarez
                (Rivermont Ledger) for police overtime records, January–June.

                Action items:
                - Alan, review the responsive records for exemptions/redactions
                  (CJIS-adjacent fields) by July 24.
                - I'll log the request in the FOIA register today.
                - Records to send the response letter + released records to the
                  requester by July 31 — that's our statutory 10-business-day deadline.

                Draft cover letter is records-cover.docx; retention rules in
                retention-schedule.xlsx.

                Sofia Marin
                City Clerk, Rivermont"""),
        },
        {
            "from": "J. Alvarez <jalvarez@rivermontledger.example>",
            "to": "records@cityofrivermont.example",
            "date": "Wed, 15 Jul 2026 18:03:00 -0500",
            "body": textwrap.dedent("""\
                Under the state public records act, I request all police department
                overtime records for January through June of this year, including
                totals by officer and by month. Electronic copies preferred.

                Thank you,
                J. Alvarez, Rivermont Ledger"""),
        },
    ],
}

MUNICIPAL_TODOS = [
    {"assignee": "Alan Woods", "task": "Review responsive records for exemptions/redactions (CJIS-adjacent)",
     "due": "2026-07-24", "due_text": "by July 24", "source": "Sofia Marin, 16 Jul 09:12"},
    {"assignee": "Sofia Marin", "task": "Log FOIA request #2026-0417 in the FOIA register",
     "due": "2026-07-16", "due_text": "today", "source": "Sofia Marin, 16 Jul 09:12"},
    {"assignee": "Mahmoud Hafez", "task": "Send response letter + released records to requester",
     "due": "2026-07-31", "due_text": "by July 31 (statutory deadline)", "source": "Sofia Marin, 16 Jul 09:12"},
]

# --- Renderers ---------------------------------------------------------------


def _p(rel: str) -> str:
    path = os.path.join(OUT, rel)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    return path


def write_eml(thread: dict) -> str:
    """RFC822-style thread: newest message with headers, older ones quoted."""
    msgs = thread["messages"]
    top = msgs[0]
    lines = [
        f"From: {top['from']}",
        f"To: {top['to']}",
        f"Date: {top['date']}",
        f"Subject: {thread['subject']}",
        "",
        top["body"],
        "",
    ]
    quote_prefix = ""
    for older in msgs[1:]:
        quote_prefix += "> "
        lines.append(f"{quote_prefix}On {older['date']}, {older['from']} wrote:")
        for bl in older["body"].splitlines():
            lines.append(f"{quote_prefix}{bl}".rstrip())
        lines.append("")
    path = _p(thread["filename"])
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines).rstrip() + "\n")
    return path


def write_docx(rel: str, title: str, paras: list[str]) -> str:
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    for para in paras:
        if para.startswith("- "):
            doc.add_paragraph(para[2:], style="List Bullet")
        else:
            doc.add_paragraph(para)
    path = _p(rel)
    doc.save(path)
    return path


def write_xlsx(rel: str, sheet: str, header: list[str], rows: list[list]) -> str:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    ws.append(header)
    for r in rows:
        ws.append(r)
    path = _p(rel)
    wb.save(path)
    return path


def write_pdf(rel: str, blocks: list[tuple[str, list[str]]]) -> str:
    """One page per (title, lines) block — used for slide-decks and record pages."""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    path = _p(rel)
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    for title, lines in blocks:
        c.setFont("Helvetica-Bold", 20)
        c.drawString(72, height - 96, title)
        c.setFont("Helvetica", 12)
        y = height - 140
        for ln in lines:
            c.drawString(72, y, ln)
            y -= 20
        c.showPage()
    c.save()
    return path


def write_json(rel_under_expected: str, obj) -> str:
    path = os.path.join(EXPECTED, rel_under_expected)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2)
        f.write("\n")
    return path


# --- Scenario assembly -------------------------------------------------------


def build_enterprise() -> list[str]:
    out = [write_eml(ENTERPRISE_THREAD)]
    out.append(write_docx(
        "enterprise/docs/spend-freeze-memo.docx",
        "Q3 Spending & Hiring Freeze — Policy Memo",
        [
            "Effective July 1, 2026 through the end of Q3, all non-critical spending "
            "and hiring is paused pending reforecast.",
            "A hire is 'critical' only if it backfills a departure in a revenue or "
            "compliance-bearing role, or is required by an executed customer contract.",
            "- Open reqs not meeting the critical bar are held, not cancelled.",
            "- Tooling and cloud spend over $5k requires CFO sign-off this quarter.",
            "Questions: People Ops (Priya Nair) or Finance (Dana Ruiz).",
        ],
    ))
    out.append(write_xlsx(
        "enterprise/sheets/q3-budget.xlsx", "Q3 Opex",
        ["Department", "Category", "Budget", "Actual", "Variance", "Variance %"],
        [
            ["Engineering", "Cloud", 180000, 214000, 34000, "18.9%"],
            ["Engineering", "Tooling", 60000, 78000, 18000, "30.0%"],
            ["Sales", "Travel", 45000, 41000, -4000, "-8.9%"],
            ["Marketing", "Events", 90000, 96000, 6000, "6.7%"],
            ["G&A", "Facilities", 70000, 72000, 2000, "2.9%"],
            ["TOTAL", "", 445000, 501000, 56000, "12.6%"],
        ],
    ))
    out.append(write_pdf(
        "enterprise/decks/q3-review.pdf",
        [
            ("Q3 Budget Review", ["Northwind — CFO review", "Board pre-read", "15 Jul 2026"]),
            ("Where we are", ["Opex tracking 12.6% over plan", "Driver: cloud + tooling",
                              "Sales travel favorable (-9%)"]),
            ("The plan", ["Cut $40k tooling (Eng)", "Pause 2 backend reqs",
                          "CFO sign-off on >$5k cloud/tooling"]),
            ("Asks", ["Revised headcount plan — Fri", "Reforecast by end of month",
                      "Board update — Monday"]),
        ],
    ))
    out.append(write_json("enterprise-todos.json", {
        "scenario": "enterprise",
        "source": ENTERPRISE_THREAD["filename"],
        "todos": ENTERPRISE_TODOS,
    }))
    return out


def build_municipal() -> list[str]:
    out = [write_eml(MUNICIPAL_THREAD)]
    out.append(write_docx(
        "municipal/docs/records-cover.docx",
        "Public Records Response — FOIA #2026-0417 (DRAFT)",
        [
            "Dear J. Alvarez,",
            "In response to your public-records request received July 15, 2026, the "
            "City of Rivermont encloses responsive police overtime records for "
            "January through June 2026.",
            "Certain fields have been redacted under the applicable CJIS and personnel "
            "exemptions; a redaction log is included with the records.",
            "- Records released: overtime totals by officer and by month.",
            "- Redacted: home addresses, badge PII, and any CJIS-restricted identifiers.",
            "Sincerely, Office of the City Clerk",
        ],
    ))
    out.append(write_xlsx(
        "municipal/sheets/retention-schedule.xlsx", "Retention",
        ["Record Series", "Retention", "Disposition", "Statutory Cite"],
        [
            ["Police overtime records", "3 years", "Destroy", "RC 149.351"],
            ["Public-records requests (FOIA log)", "2 years", "Archive", "RC 149.43"],
            ["Personnel files", "Term + 6 years", "Restricted", "RC 149.011"],
            ["CJIS-restricted data", "Per CJIS policy", "Do not disclose", "CJIS 5.9"],
        ],
    ))
    out.append(write_pdf(
        "municipal/records/responsive-records.pdf",
        [
            ("Police Overtime — Jan–Jun 2026", ["City of Rivermont", "FOIA #2026-0417",
                                                "Released records (redacted)"]),
            ("Overtime by month", ["Jan: 412 hrs / $28,140", "Feb: 388 hrs / $26,900",
                                   "Mar: 455 hrs / $31,200", "Apr: 401 hrs / $27,500",
                                   "May: 470 hrs / $32,800", "Jun: 498 hrs / $34,600"]),
            ("Top officers (badge redacted)", ["Officer [REDACTED] — 210 hrs",
                                               "Officer [REDACTED] — 188 hrs",
                                               "Officer [REDACTED] — 176 hrs"]),
        ],
    ))
    out.append(write_json("municipal-todos.json", {
        "scenario": "municipal",
        "source": MUNICIPAL_THREAD["filename"],
        "todos": MUNICIPAL_TODOS,
    }))
    return out


def main():
    written = build_enterprise() + build_municipal()
    print(f"Wrote {len(written)} files:")
    for p in written:
        print("  ", os.path.relpath(p, HERE))


if __name__ == "__main__":
    main()
